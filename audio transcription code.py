import torch
import os
from datetime import datetime
import threading
import time
import librosa
from docx import Document
import json
import customtkinter as ctk
from tkinter import filedialog, messagebox
import tkinter as tk
from PIL import Image, ImageTk
import sys
import platform

# Try to import accelerate - it's optional but will improve performance
try:
    import accelerate
    HAS_ACCELERATE = True
except ImportError:
    HAS_ACCELERATE = False

# Now import the transformers components
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline

# Définir les thèmes et l'apparence
ctk.set_appearance_mode("System")  # Modes: "Dark", "Light", "System"
ctk.set_default_color_theme("blue")  # Thèmes: "blue", "green", "dark-blue"

class ModernAudioTranscriptionApp(ctk.CTk):
    def __init__(self, root=None):
        super().__init__()
        
        # Configuration de la fenêtre
        self.title("AudioTrans Pro")
        self.geometry("950x680")
        self.minsize(900, 650)

        # Variables de l'application
        self.file_path = ctk.StringVar()
        self.language = ctk.StringVar(value="fr")
        self.output_format = ctk.StringVar(value="txt")
        self.status_var = ctk.StringVar(value="Prêt")
        self.device_info = ctk.StringVar(value="Périphérique: Non détecté")
        self.progress_value = ctk.DoubleVar(value=0)
        self.progress_text = ctk.StringVar(value="0%")
        self.theme_var = ctk.StringVar(value=ctk.get_appearance_mode())
        # Nouvelle variable pour la précision
        self.precision_value = ctk.DoubleVar(value=0.0)
        self.precision_text = ctk.StringVar(value="Précision: Élevée")
        # Nouveaux paramètres
        self.timestamps_var = ctk.BooleanVar(value=False)
        self.beam_size_var = ctk.IntVar(value=1)

        # Mappings
        self.lang_mapping = {
            "Français": "fr", 
            "Anglais": "en", 
            "Allemand": "de", 
            "Espagnol": "es", 
            "Italien": "it"
        }
        self.reverse_lang_mapping = {v: k for k, v in self.lang_mapping.items()}
        self.format_mapping = {"Texte (.txt)": "txt", "Document Word (.docx)": "docx"}
        self.reverse_format_mapping = {v: k for k, v in self.format_mapping.items()}
        
        # Variables du modèle
        self.model = None
        self.processor = None
        self.pipe = None
        self.audio_duration = 0
        self.progress_stop = False
        self.transcription_running = False  # Nouvelle variable pour suivre l'état de la transcription
        
        # Charger les préférences si disponibles
        self.load_preferences()
        
        # Configurer l'interface utilisateur
        self.setup_ui()
        
    def load_preferences(self):
        """Charger les préférences utilisateur depuis un fichier JSON"""
        try:
            if os.path.exists("audiotrans_preferences.json"):
                with open("audiotrans_preferences.json", "r") as f:
                    prefs = json.load(f)
                    
                    if "theme" in prefs:
                        ctk.set_appearance_mode(prefs["theme"])
                        self.theme_var.set(prefs["theme"])
                    
                    if "language" in prefs:
                        self.language.set(prefs["language"])
                    
                    if "format" in prefs:
                        self.output_format.set(prefs["format"])
                        
                    # Chargement de la valeur de précision
                    if "precision" in prefs:
                        self.precision_value.set(prefs["precision"])
                        self.update_precision_label(prefs["precision"])
                        
                    # Chargement des nouveaux paramètres
                    if "timestamps" in prefs:
                        self.timestamps_var.set(prefs["timestamps"])
                    if "beam_size" in prefs:
                        self.beam_size_var.set(prefs["beam_size"])
        except Exception as e:
            print(f"Erreur lors du chargement des préférences: {e}")
            
    def save_preferences(self):
        """Sauvegarder les préférences utilisateur dans un fichier JSON"""
        try:
            prefs = {
                "theme": self.theme_var.get(),
                "language": self.language.get(),
                "format": self.output_format.get(),
                "precision": self.precision_value.get(),  # Sauvegarde de la précision
                "timestamps": self.timestamps_var.get(),  # Sauvegarde des timestamps
                "beam_size": self.beam_size_var.get()     # Sauvegarde de la taille du faisceau
            }
            
            with open("audiotrans_preferences.json", "w") as f:
                json.dump(prefs, f)
        except Exception as e:
            print(f"Erreur lors de la sauvegarde des préférences: {e}")
        
    def setup_ui(self):
        """Créer l'interface utilisateur moderne avec défilement"""
        # Conteneur externe pour contenir le cadre défilable
        outer_frame = ctk.CTkFrame(self, corner_radius=15)
        outer_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Remplacer le conteneur principal par un cadre défilable
        main_frame = ctk.CTkScrollableFrame(
            outer_frame, 
            corner_radius=15,
            fg_color=("gray95", "gray10"),  # Couleurs pour modes clair/sombre
        )
        main_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        # Barre d'outils supérieure avec style amélioré
        toolbar_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        toolbar_frame.pack(fill="x", padx=10, pady=(5, 15))
        
        # Logo ou titre de l'application
        app_title = ctk.CTkLabel(
            toolbar_frame, 
            text="AudioTrans Pro",
            font=("Segoe UI", 20, "bold"),
            text_color=("#1E5CAA", "#3B8ED0")  # Bleu foncé/clair pour modes sombre/clair
        )
        app_title.pack(side="left", padx=5)
        
        # Bouton de chargement de modèle
        self.load_model_btn = ctk.CTkButton(
            toolbar_frame, 
            text="Charger Modèle", 
            command=self.load_model,
            width=150,
            height=36,
            corner_radius=8,
            fg_color=("#2a6099", "#3B8ED0"),
            hover_color=("#1f4a7a", "#2980b9"),
            font=("Segoe UI", 12, "bold")
        )
        self.load_model_btn.pack(side="left", padx=15)
        
        # Information sur le périphérique
        device_label = ctk.CTkLabel(
            toolbar_frame, 
            textvariable=self.device_info,
            font=("Segoe UI", 12)
        )
        device_label.pack(side="left", padx=10)
        
        # Bouton de basculement de thème
        theme_btn = ctk.CTkButton(
            toolbar_frame, 
            text="Changer Thème", 
            command=self.toggle_theme,
            width=130,
            height=36,
            corner_radius=8,
            font=("Segoe UI", 12)
        )
        theme_btn.pack(side="right", padx=5)
        
        # Sélection de fichier audio avec apparence améliorée
        file_frame = ctk.CTkFrame(main_frame, corner_radius=10, fg_color=("gray95", "gray20"))
        file_frame.pack(fill="x", padx=10, pady=(0, 15))
        
        file_label = ctk.CTkLabel(
            file_frame, 
            text="Fichier Audio:", 
            font=("Segoe UI", 14, "bold")
        )
        file_label.pack(side="left", padx=(15, 10), pady=15)
        
        file_entry_frame = ctk.CTkFrame(file_frame, fg_color="transparent")
        file_entry_frame.pack(side="left", fill="x", expand=True, padx=(0, 15), pady=15)
        
        file_entry = ctk.CTkEntry(
            file_entry_frame, 
            textvariable=self.file_path, 
            width=400,
            height=36,
            corner_radius=8,
            font=("Segoe UI", 12)
        )
        file_entry.pack(side="left", fill="x", expand=True)
        
        # Bouton parcourir avec style amélioré
        self.browse_button = ctk.CTkButton(
            file_entry_frame, 
            text="Parcourir",
            command=self.browse_file, 
            width=120,
            height=36,
            corner_radius=8,
            font=("Segoe UI", 12, "bold")
        )
        self.browse_button.pack(side="right", padx=(15, 0))
        
        # Cadre pour les options avec style amélioré
        options_frame = ctk.CTkFrame(main_frame, corner_radius=10, fg_color=("gray95", "gray20"))
        options_frame.pack(fill="x", padx=10, pady=(0, 15))
        
        # Section des options avec meilleure organisation
        options_inner_frame = ctk.CTkFrame(options_frame, fg_color="transparent")
        options_inner_frame.pack(fill="x", padx=15, pady=15)
        
        # Grille pour les options
        options_inner_frame.grid_columnconfigure(0, weight=0)
        options_inner_frame.grid_columnconfigure(1, weight=1)
        options_inner_frame.grid_columnconfigure(2, weight=0)
        options_inner_frame.grid_columnconfigure(3, weight=1)
        
        # Langue
        lang_label = ctk.CTkLabel(
            options_inner_frame, 
            text="Langue:", 
            font=("Segoe UI", 13, "bold")
        )
        lang_label.grid(row=0, column=0, sticky="w", padx=(0, 10), pady=5)
        
        lang_menu = ctk.CTkOptionMenu(
            options_inner_frame,
            values=list(self.lang_mapping.keys()),
            command=self.on_lang_select,
            width=180,
            height=36,
            dynamic_resizing=False,
            font=("Segoe UI", 12)
        )
        lang_menu.grid(row=0, column=1, sticky="w", padx=5, pady=5)
        
        # Format
        format_label = ctk.CTkLabel(
            options_inner_frame, 
            text="Format:", 
            font=("Segoe UI", 13, "bold")
        )
        format_label.grid(row=0, column=2, sticky="w", padx=(20, 10), pady=5)
        
        format_menu = ctk.CTkOptionMenu(
            options_inner_frame,
            values=list(self.format_mapping.keys()),
            command=self.on_format_select,
            width=180,
            height=36,
            dynamic_resizing=False,
            font=("Segoe UI", 12)
        )
        format_menu.grid(row=0, column=3, sticky="w", padx=5, pady=5)
        
        # Nouveau: curseur de précision de transcription
        precision_frame = ctk.CTkFrame(options_inner_frame, fg_color="transparent")
        precision_frame.grid(row=1, column=0, columnspan=4, sticky="ew", padx=0, pady=(15, 0))
        
        precision_label = ctk.CTkLabel(
            precision_frame, 
            text="Niveau de précision:", 
            font=("Segoe UI", 13, "bold")
        )
        precision_label.pack(side="left", padx=(0, 10))
        
        precision_value_label = ctk.CTkLabel(
            precision_frame, 
            textvariable=self.precision_text,
            font=("Segoe UI", 12)
        )
        precision_value_label.pack(side="right", padx=10)
        
        precision_slider = ctk.CTkSlider(
            options_inner_frame,
            from_=0.0,
            to=1.0,
            number_of_steps=20,
            variable=self.precision_value,
            command=self.on_precision_change,
            width=550,
            height=20,
            border_width=5,
            progress_color=("#2a6099", "#3B8ED0")
        )
        precision_slider.grid(row=2, column=0, columnspan=4, sticky="ew", padx=0, pady=(5, 10))
        
        # Options avancées avec cadre dédié
        advanced_frame = ctk.CTkFrame(main_frame, corner_radius=10, fg_color=("gray95", "gray20"))
        advanced_frame.pack(fill="x", padx=10, pady=(0, 15))
        
        advanced_label = ctk.CTkLabel(
            advanced_frame, 
            text="Options avancées", 
            font=("Segoe UI", 14, "bold")
        )
        advanced_label.pack(anchor="w", padx=15, pady=(15, 10))
        
        advanced_inner_frame = ctk.CTkFrame(advanced_frame, fg_color="transparent")
        advanced_inner_frame.pack(fill="x", padx=15, pady=(0, 15))
        
        # Case à cocher pour les horodatages
        timestamps_cb = ctk.CTkCheckBox(
            advanced_inner_frame,
            text="Inclure les horodatages",
            variable=self.timestamps_var,
            font=("Segoe UI", 12),
            command=self.save_preferences
        )
        timestamps_cb.pack(side="left", padx=(0, 30))
        
        # Cadre pour la qualité de transcription
        beam_frame = ctk.CTkFrame(advanced_inner_frame, fg_color="transparent")
        beam_frame.pack(side="left")
        
        beam_label = ctk.CTkLabel(
            beam_frame,
            text="Qualité de transcription:",
            font=("Segoe UI", 12)
        )
        beam_label.pack(side="left", padx=(0, 10))
        
        beam_options = ctk.CTkSegmentedButton(
            beam_frame,
            values=["Rapide", "Standard", "Élevée"],
            command=self.on_beam_change,
            font=("Segoe UI", 12)
        )
        beam_options.pack(side="left")
        
        # Définir la valeur initiale
        if self.beam_size_var.get() == 1:
            beam_options.set("Rapide")
        elif self.beam_size_var.get() == 2:  # Modifié: 3 -> 2
            beam_options.set("Standard")
        else:
            beam_options.set("Élevée")
        
        # Initialisation des menus déroulants
        lang_menu.set(self.reverse_lang_mapping.get(self.language.get(), "Français"))
        format_menu.set(self.reverse_format_mapping.get(self.output_format.get(), "Texte (.txt)"))
        
        # Zone de résultat avec style amélioré
        result_frame = ctk.CTkFrame(main_frame, corner_radius=10)
        result_frame.pack(fill="both", expand=True, padx=10, pady=(0, 15))
        
        result_label = ctk.CTkLabel(
            result_frame, 
            text="Résultat de la transcription", 
            font=("Segoe UI", 14, "bold")
        )
        result_label.pack(anchor="w", padx=15, pady=(15, 5))
        
        # Zone de texte pour les résultats avec style amélioré
        self.result_text = ctk.CTkTextbox(
            result_frame, 
            wrap="word",
            font=("Segoe UI", 12),
            corner_radius=8,
            border_width=2,
            border_color=("gray80", "gray30"),
            height=200
        )
        self.result_text.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        
        # Barre de statut et progression avec style amélioré
        status_frame = ctk.CTkFrame(main_frame, corner_radius=10, fg_color=("gray95", "gray20"))
        status_frame.pack(fill="x", padx=10, pady=(0, 0))
        
        # Section supérieure de la barre de statut
        status_top_frame = ctk.CTkFrame(status_frame, fg_color="transparent")
        status_top_frame.pack(fill="x", padx=15, pady=(15, 5))
        
        # Statut
        status_label = ctk.CTkLabel(
            status_top_frame, 
            text="Statut:", 
            font=("Segoe UI", 13, "bold")
        )
        status_label.pack(side="left")
        
        status_text = ctk.CTkLabel(
            status_top_frame, 
            textvariable=self.status_var,
            font=("Segoe UI", 12)
        )
        status_text.pack(side="left", padx=(5, 20))
        
        # Temps
        self.time_label = ctk.CTkLabel(
            status_top_frame, 
            text="Temps: --",
            font=("Segoe UI", 12)
        )
        self.time_label.pack(side="left", padx=10)
        
        # Progression et texte de progression
        progress_frame = ctk.CTkFrame(status_frame, fg_color="transparent")
        progress_frame.pack(fill="x", padx=15, pady=(0, 15))
        
        self.progress_bar = ctk.CTkProgressBar(
            progress_frame,
            variable=self.progress_value,
            width=700,
            height=15,
            corner_radius=5,
            border_width=0,
            progress_color=("#2a6099", "#3B8ED0")
        )
        self.progress_bar.pack(side="left", fill="x", expand=True, padx=(0, 10))
        
        progress_text_label = ctk.CTkLabel(
            progress_frame, 
            textvariable=self.progress_text,
            font=("Segoe UI", 12),
            width=50
        )
        progress_text_label.pack(side="right")
        
        # Boutons d'action avec style amélioré
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(fill="x", padx=10, pady=(15, 0))
        
        # Bouton de transcription
        self.transcribe_btn = ctk.CTkButton(
            button_frame, 
            text="Transcrire", 
            command=self.start_transcription,
            width=170,
            height=40,
            corner_radius=8,
            font=("Segoe UI", 14, "bold"),
            fg_color=("#2a6099", "#3B8ED0"),
            hover_color=("#1f4a7a", "#2980b9")
        )
        self.transcribe_btn.pack(side="left", padx=5)
        
        # Nouveau: Bouton d'annulation
        self.cancel_btn = ctk.CTkButton(
            button_frame, 
            text="Annuler", 
            command=self.cancel_transcription,
            width=130,
            height=40,
            corner_radius=8,
            font=("Segoe UI", 14, "bold"),
            fg_color=("#d63031", "#e74c3c"),
            hover_color=("#b71c1c", "#c0392b"),
            state="disabled"
        )
        self.cancel_btn.pack(side="left", padx=5)
        
        # Bouton de sauvegarde
        self.save_btn = ctk.CTkButton(
            button_frame, 
            text="Sauvegarder", 
            command=self.save_transcription,
            width=170,
            height=40,
            corner_radius=8,
            state="disabled",
            font=("Segoe UI", 14, "bold")
        )
        self.save_btn.pack(side="left", padx=5)
        
        # Initialisation de l'UI
        self.result_text.insert("1.0", "Bienvenue dans AudioTrans Pro!\n\n"
                               "Pour commencer:\n"
                               "1. Chargez le modèle\n"
                               "2. Sélectionnez un fichier audio\n"
                               "3. Choisissez la langue et le format\n"
                               "4. Cliquez sur Transcrire\n\n"
                               "Le curseur de précision vous permet d'ajuster la qualité de transcription:\n"
                               "- Gauche: Transcription plus précise mais moins créative\n"
                               "- Droite: Transcription plus créative mais potentiellement moins précise\n")
        self.result_text.configure(state="disabled")
        
    def toggle_theme(self):
        """Basculer entre les thèmes clair et sombre"""
        current_theme = self.theme_var.get()
        new_theme = "Light" if current_theme == "Dark" else "Dark"
        self.theme_var.set(new_theme)
        ctk.set_appearance_mode(new_theme)
        self.save_preferences()
        
    def on_lang_select(self, choice):
        """Gestion de la sélection de langue"""
        self.language.set(self.lang_mapping[choice])
        self.save_preferences()
        
    def on_format_select(self, choice):
        """Gestion de la sélection de format"""
        self.output_format.set(self.format_mapping[choice])
        self.save_preferences()
        
    def on_precision_change(self, value):
        """Callback pour le changement de valeur du curseur de précision"""
        self.update_precision_label(float(value))
        # Sauvegarde automatique des préférences lors du changement
        self.save_preferences()
        
    def update_precision_label(self, value):
        """Met à jour le texte de l'étiquette de précision basé sur la valeur"""
        # Inversé: 0.0 = très précis, 1.0 = plus créatif
        if value <= 0.2:
            self.precision_text.set("Précision: Très élevée")
        elif value <= 0.4:
            self.precision_text.set("Précision: Élevée")
        elif value <= 0.6:
            self.precision_text.set("Précision: Moyenne")
        elif value <= 0.8:
            self.precision_text.set("Précision: Basse")
        else:
            self.precision_text.set("Précision: Créative")
        
    def browse_file(self):
        """Ouvrir la boîte de dialogue pour sélectionner un fichier audio"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Fichiers Audio", "*.mp3 *.wav *.flac *.ogg"), ("Tous les fichiers", "*.*")]
        )
        if file_path:
            self.file_path.set(file_path)
            self.get_audio_duration(file_path)
            
    def get_audio_duration(self, file_path):
        """Obtenir la durée du fichier audio"""
        # Vérifier si le chemin du fichier est valide
        if not file_path or not os.path.exists(file_path):
            self.audio_duration = 0
            self.status_var.set("Fichier audio invalide ou inexistant")
            return 0
            
        try:
            # Utiliser librosa pour obtenir la durée audio
            self.audio_duration = librosa.get_duration(path=file_path)
            # S'assurer que la durée n'est pas None
            if self.audio_duration is None:
                self.audio_duration = 0
                
            duration_min = int(self.audio_duration // 60)
            duration_sec = int(self.audio_duration % 60)
            self.status_var.set(f"Prêt - Durée audio: {duration_min}m {duration_sec}s")
            return self.audio_duration
        except Exception as e:
            self.audio_duration = 0
            self.status_var.set(f"Impossible de déterminer la durée audio: {str(e)}")
            return 0
            
    def update_result_text(self, text, append=False):
        """Mettre à jour le texte de résultat"""
        self.result_text.configure(state="normal")
        
        if not append:
            self.result_text.delete("1.0", "end")
        
        # Si le résultat contient des timestamps (c'est un dictionnaire avec "chunks")
        if isinstance(text, dict) and "chunks" in text:
            formatted_text = "Transcription avec horodatages :\n\n"
            for chunk in text["chunks"]:
                start = chunk.get("timestamp", [0])[0]
                # Formater le temps en minutes:secondes
                minutes = int(start // 60)
                seconds = int(start % 60)
                time_str = f"[{minutes:02d}:{seconds:02d}] "
                chunk_text = chunk.get("text", "")
                formatted_text += time_str + chunk_text + "\n"
            self.result_text.insert("end", formatted_text)
        else:
            # Affichage normal du texte
            self.result_text.insert("end", text)
        
        self.result_text.configure(state="disabled")
        self.result_text.see("end")
        
    def load_model(self):
        """Charger le modèle Whisper"""
        def _load():
            try:
                self.status_var.set("Chargement du modèle...")
                self.progress_value.set(0)
                self.progress_text.set("Chargement...")
                self.load_model_btn.configure(state="disabled")
                
                # Configuration du périphérique
                device = "cuda:0" if torch.cuda.is_available() else "cpu"
                torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32
                
                # Mise à jour des informations du périphérique
                device_name = "GPU" if torch.cuda.is_available() else "CPU"
                self.device_info.set(f"Périphérique: {device_name} ({device})")
                
                # Chargement du modèle et du processeur
                model_id = "openai/whisper-large-v3"
                
                self.status_var.set("Chargement du modèle Whisper...")
                self.progress_value.set(0.2)
                self.progress_text.set("20%")
                time.sleep(0.5)  # Permettre à l'UI de se mettre à jour
                
                # Optimisations pour la mémoire et les performances
                # Vérifier si accélerate est disponible avant d'utiliser les options qui en dépendent
                model_kwargs = {
                    "torch_dtype": torch_dtype,
                    "attn_implementation": "eager"  # Optimisation pour l'attention
                }
                
                # Ajouter les options qui nécessitent Accelerate uniquement si disponible
                if HAS_ACCELERATE:
                    model_kwargs.update({
                        "low_cpu_mem_usage": True,
                        "use_safetensors": True
                    })
                
                self.model = AutoModelForSpeechSeq2Seq.from_pretrained(
                    model_id,
                    **model_kwargs
                )
                
                self.progress_value.set(0.6)
                self.progress_text.set("60%")
                self.status_var.set("Chargement du processeur...")
                time.sleep(0.5)  # Permettre à l'UI de se mettre à jour
                
                self.model.to(device)
                self.processor = AutoProcessor.from_pretrained(model_id)
                
                self.progress_value.set(0.8)
                self.progress_text.set("80%")
                self.status_var.set("Initialisation du pipeline...")
                time.sleep(0.5)  # Permettre à l'UI de se mettre à jour
                
                # Création du pipeline avec des paramètres optimisés
                self.pipe = pipeline(
                    "automatic-speech-recognition",
                    model=self.model,
                    tokenizer=self.processor.tokenizer,
                    feature_extractor=self.processor.feature_extractor,
                    chunk_length_s=30,
                    stride_length_s=5,
                    torch_dtype=torch_dtype,
                    device=device
                )
                
                self.progress_value.set(1.0)
                self.progress_text.set("100%")
                self.status_var.set("Modèle chargé avec succès - Prêt à transcrire")
                self.transcribe_btn.configure(state="normal")
                
                # Information rapide dans la zone de résultats
                self.update_result_text("Modèle chargé avec succès.\n\nSélectionnez un fichier audio et cliquez sur 'Transcrire' pour commencer.")
                
            except Exception as e:
                self.status_var.set(f"Erreur lors du chargement du modèle")
                self.update_result_text(f"Erreur lors du chargement du modèle: {str(e)}\n")
                self.progress_value.set(0)
                self.progress_text.set("Erreur")
            finally:
                self.load_model_btn.configure(state="normal")
                
        # Exécuter dans un thread séparé pour éviter le gel de l'interface
        threading.Thread(target=_load, daemon=True).start()

    # Ajout d'une propriété pour stocker la dernière transcription
    @property
    def transcription_result(self):
        return self._transcription_result if hasattr(self, '_transcription_result') else ""
        
    @transcription_result.setter
    def transcription_result(self, value):
        self._transcription_result = value
        
    def simulate_progress(self, duration, beam_size=1):
        """Simuler la progression de la transcription en tenant compte de la complexité"""
        self.progress_value.set(0)
        self.progress_text.set("0%")
        self.progress_stop = False
        
        # S'assurer que duration n'est pas None
        if duration is None or duration <= 0:
            duration = 30  # valeur par défaut si la durée est inconnue ou invalide
            
        # Estimer le temps total (la transcription est généralement plus rapide que le temps réel)
        # Ajuster le temps estimé en fonction de la complexité (beam_size)
        complexity_factor = 1.0
        if beam_size == 2:  # Standard (réduit de 3 à 2)
            complexity_factor = 2.0
        elif beam_size == 3:  # Élevée (réduit de 5 à 3)
            complexity_factor = 3.0
        
        # Nouvelle approche: diviser la progression en phases
        total_steps = 100
        
        if beam_size > 1:  # Pour Standard et Élevée
            # Phases de progression pour les modes avancés
            # Phase 1: 0-80% - progression normale
            # Phase 2: 80-98% - progression plus lente (changé de 95% à 98%)
            # Reste: attendre la fin réelle de la transcription
            phase1_pct = 80   # Jusqu'à où va la phase 1 (augmenté de 75% à 80%)
            phase2_pct = 98   # Jusqu'à où va la phase 2 (augmenté de 95% à 98%)
            
            phase1_steps = int(total_steps * phase1_pct / 100)
            phase2_steps = int(total_steps * phase2_pct / 100)
        else:
            # Progression standard pour le mode rapide
            phase1_pct = 90
            phase2_pct = 98
            
            phase1_steps = int(total_steps * phase1_pct / 100)
            phase2_steps = int(total_steps * phase2_pct / 100)
        
        # Calcul des temps pour chaque phase
        estimated_time = max(3, duration * 0.5 * complexity_factor)
        phase1_time = estimated_time * 0.6  # 60% du temps pour la première phase
        phase2_time = estimated_time * 0.4  # 40% du temps pour la deuxième phase
        
        # Intervalle entre les mises à jour en ms pour chaque phase
        phase1_interval = int((phase1_time * 1000) / phase1_steps)
        phase2_interval = int((phase2_time * 1000) / (phase2_steps - phase1_steps))
        
        self.progress_thread = threading.Thread(target=self._update_progress, 
                                             args=(phase1_steps, phase2_steps, phase1_interval, phase2_interval))
        self.progress_thread.daemon = True
        self.progress_thread.start()
    
    def _update_progress(self, phase1_steps, phase2_steps, phase1_interval, phase2_interval):
        """Mise à jour de la progression en plusieurs phases - CORRIGÉ"""
        step = 0
        total_steps = 100
        
        # Phase 1: progression rapide
        while step < phase1_steps and not self.progress_stop:
            step += 1
            progress = step / total_steps
            
            # Mettre à jour l'interface
            self.progress_value.set(progress)
            self.progress_text.set(f"{int(progress * 100)}%")
            
            # Pause entre les mises à jour
            time.sleep(phase1_interval / 1000)
        
        # Phase 2: progression plus lente
        while step < phase2_steps and not self.progress_stop:
            step += 1
            progress = step / total_steps
            
            # Mettre à jour l'interface
            self.progress_value.set(progress)
            self.progress_text.set(f"{int(progress * 100)}%")
            
            # Pause plus longue entre les mises à jour
            time.sleep(phase2_interval / 1000)
        
        # Phase 3: rester à 99% en attendant la fin réelle
        if not self.progress_stop:
            self.progress_value.set(0.99)
            self.progress_text.set("99%")
            
        # La méthode complete_progress sera appelée pour finaliser à 100%
        
    def stop_progress(self):
        """Arrêter la simulation de progression"""
        self.progress_stop = True
        
    def complete_progress(self):
        """Finaliser la progression à 100% - CORRIGÉ"""
        # D'abord arrêter la simulation
        self.progress_stop = True
        
        # Petit délai pour s'assurer que la simulation est bien arrêtée
        time.sleep(0.1)
        
        # Forcer la mise à jour directement, sans délai
        self.progress_value.set(1.0)
        self.progress_text.set("100%")
        self.progress_bar.set(1.0)
        
        # Forcer le rafraîchissement de l'interface
        self.update_idletasks()
        
        # Afficher le message de complétion
        self.status_var.set("Transcription terminée!")
        
    def cancel_transcription(self):
        """Annuler la transcription en cours"""
        if self.transcription_running:
            self.progress_stop = True
            self.transcription_running = False
            self.status_var.set("Transcription annulée")
            self.progress_value.set(0)
            self.progress_text.set("0%")
            self.cancel_btn.configure(state="disabled")
            self.transcribe_btn.configure(state="normal")
            self.update_result_text("La transcription a été annulée par l'utilisateur.")
            
    def start_transcription(self):
        """Démarrer le processus de transcription"""
        # Vérifier qu'un fichier audio a été sélectionné
        audio_file = self.file_path.get()
        if not audio_file or not os.path.exists(audio_file):
            messagebox.showerror("Erreur", "Veuillez sélectionner un fichier audio valide")
            self.status_var.set("Erreur: fichier audio non valide")
            return
            
        # Vérifier que le modèle est chargé
        if self.pipe is None:
            messagebox.showinfo("Information", "Chargement du modèle requis", 
                              detail="Veuillez d'abord charger le modèle de transcription.")
            self.status_var.set("Veuillez charger le modèle")
            return
            
        # Obtenir la durée de l'audio
        self.audio_duration = self.get_audio_duration(audio_file)
        
        # Récupérer la taille du faisceau (beam_size) pour l'estimation de la progression
        beam_size = self.beam_size_var.get()
            
        def _transcribe():
            # Variable pour indiquer qu'une transcription est en cours
            self.transcription_running = True
            
            try:
                self.transcribe_btn.configure(state="disabled")
                self.cancel_btn.configure(state="normal")  # Activer le bouton d'annulation
                self.status_var.set("Transcription en cours...")
                self.update_result_text("Démarrage de la transcription...\n")
                
                # Démarrer la simulation de progression avec le beam_size pour une meilleure estimation
                self.simulate_progress(self.audio_duration, beam_size)
                
                start_time = time.time()
                selected_lang = self.language.get()
                
                # Obtenir la valeur de précision (température)
                temperature = self.precision_value.get()
                
                # Obtenir les valeurs des nouveaux paramètres
                use_timestamps = self.timestamps_var.get()
                
                # Variables pour le résultat et les erreurs
                result = [None]
                error = [None]
                completed = [False]
                
                # Fonction pour exécuter la transcription
                def run_transcription():
                    try:
                        # Effectuer la transcription avec les nouveaux paramètres
                        result[0] = self.pipe(
                            audio_file,
                            batch_size=4,
                            return_timestamps=use_timestamps,
                            generate_kwargs={
                                "task": "transcribe",
                                "language": selected_lang,
                                "temperature": temperature,
                                "num_beams": beam_size
                            }
                        )
                        completed[0] = True
                    except Exception as e:
                        error[0] = e
                
                # Démarrer la transcription dans un thread
                trans_thread = threading.Thread(target=run_transcription, daemon=True)
                trans_thread.start()
                
                # Calculer le timeout en fonction de la durée audio et du beam_size
                # Plus le beam_size est grand, plus le timeout est long
                timeout = max(60, self.audio_duration * (2 + beam_size))
                
                # Attendre que la transcription se termine ou le timeout
                trans_thread.join(timeout)
                
                # Vérifier si la transcription a été annulée
                if self.progress_stop and not completed[0]:
                    # La transcription a été annulée, on sort proprement
                    self.transcription_running = False
                    return
                
                # Vérifier si la transcription s'est terminée avant le timeout
                if not completed[0]:
                    if error[0]:
                        raise error[0]
                    else:
                        raise TimeoutError(f"La transcription a pris trop de temps (>{timeout}s)")
                
                # Transcription terminée avec succès
                end_time = time.time()
                time_taken = end_time - start_time
                
                # Formatage du temps pris
                if time_taken < 60:
                    time_str = f"{time_taken:.1f} secondes"
                else:
                    minutes = int(time_taken // 60)
                    seconds = int(time_taken % 60)
                    time_str = f"{minutes} minutes {seconds} secondes"
                
                # Mettre à jour le label de temps
                self.time_label.configure(text=f"Temps: {time_str}")
                
                # Forcer la mise à jour de la barre de progression à 100%
                self.complete_progress()
                
                # Traitement du résultat selon son format (avec ou sans timestamps)
                if isinstance(result[0], dict):
                    if "chunks" in result[0]:
                        # Résultat avec timestamps
                        self.transcription_result = result[0]
                        # On utilise le texte brut pour la sauvegarde
                        raw_text = " ".join([chunk.get("text", "") for chunk in result[0]["chunks"]])
                        # Stocker le texte brut pour la sauvegarde
                        result[0]["text"] = raw_text
                    else:
                        # Résultat normal au format dict
                        self.transcription_result = result[0]["text"]
                else:
                    # Résultat sous forme de chaîne
                    self.transcription_result = str(result[0])
                    
                # Stocker les métadonnées pour la sauvegarde
                self.transcription_metadata = {
                    "source_file": audio_file,
                    "date": datetime.now(),
                    "language": selected_lang,
                    "duration": self.audio_duration,
                    "processing_time": time_taken,
                    "has_timestamps": use_timestamps
                }
                
                # Afficher la transcription dans l'affichage des résultats
                self.update_result_text(result[0] if use_timestamps else self.transcription_result)
                
                # Activer le bouton de sauvegarde
                self.save_btn.configure(state="normal")
                
            except TimeoutError as e:
                # Cas spécifique pour les timeouts - suggérer de réduire la qualité
                self.stop_progress()
                self.progress_value.set(0)
                self.progress_text.set("Erreur")
                self.status_var.set("Erreur: Timeout de transcription")
                self.update_result_text(
                    f"La transcription a pris trop de temps et a été interrompue.\n\n"
                    f"Suggestions:\n"
                    f"1. Essayez le mode 'Rapide' pour les fichiers plus longs\n"
                    f"2. Vérifiez les ressources système disponibles\n"
                    f"3. Divisez les fichiers audio très longs en segments plus courts\n"
                )
            except Exception as e:
                # Erreurs génériques
                self.stop_progress()
                self.progress_value.set(0)
                self.progress_text.set("Erreur")
                self.status_var.set(f"Erreur pendant la transcription")
                self.update_result_text(f"Erreur pendant la transcription: {str(e)}\n")
            finally:
                self.transcription_running = False
                self.transcribe_btn.configure(state="normal")
                self.cancel_btn.configure(state="disabled")  # Désactiver le bouton d'annulation
                
        # Exécuter dans un thread séparé pour éviter le gel de l'interface
        threading.Thread(target=_transcribe, daemon=True).start()
    
    def save_txt_file(self, save_path):
        """Sauvegarder la transcription sous forme de fichier texte"""
        try:
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write("TRANSCRIPTION AUDIO\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"Fichier source: {self.transcription_metadata['source_file']}\n")
                f.write(f"Date de transcription: {self.transcription_metadata['date'].strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Langue: {self.transcription_metadata['language']}\n")
                f.write(f"Durée: {self.format_duration(self.transcription_metadata['duration'])}\n")
                f.write(f"Temps de traitement: {self.format_duration(self.transcription_metadata['processing_time'])}\n\n")
                f.write("=" * 50 + "\n\n")
                
                # Si le résultat contient des timestamps, formatez-les
                has_timestamps = self.transcription_metadata.get('has_timestamps', False)
                if has_timestamps and isinstance(self.transcription_result, dict) and "chunks" in self.transcription_result:
                    f.write("TRANSCRIPTION AVEC HORODATAGES :\n\n")
                    for chunk in self.transcription_result["chunks"]:
                        start = chunk.get("timestamp", [0])[0]
                        # Formater le temps en minutes:secondes
                        minutes = int(start // 60)
                        seconds = int(start % 60)
                        time_str = f"[{minutes:02d}:{seconds:02d}] "
                        chunk_text = chunk.get("text", "")
                        f.write(time_str + chunk_text + "\n")
                else:
                    # Texte normal sans horodatage
                    if isinstance(self.transcription_result, dict) and "text" in self.transcription_result:
                        f.write(self.transcription_result["text"])
                    else:
                        f.write(str(self.transcription_result))
                
            self.status_var.set(f"Transcription sauvegardée: {save_path}")
            return True
        except Exception as e:
            self.status_var.set(f"Erreur lors de la sauvegarde: {str(e)}")
            messagebox.showerror("Erreur", f"Erreur lors de la sauvegarde: {str(e)}")
            return False
    
    def save_docx_file(self, save_path):
        """Sauvegarder la transcription sous forme de document Word"""
        try:
            # Créer un nouveau document
            doc = Document()
            
            # Ajouter un titre
            doc.add_heading('Transcription Audio', 0)
            
            # Ajouter les métadonnées
            metadata_table = doc.add_table(rows=5, cols=2)
            metadata_table.style = 'Table Grid'
            
            # Remplir les métadonnées
            cells = metadata_table.rows[0].cells
            cells[0].text = 'Fichier source'
            cells[1].text = self.transcription_metadata['source_file']
            
            cells = metadata_table.rows[1].cells
            cells[0].text = 'Date de transcription'
            cells[1].text = self.transcription_metadata['date'].strftime('%Y-%m-%d %H:%M:%S')
            
            cells = metadata_table.rows[2].cells
            cells[0].text = 'Langue'
            cells[1].text = self.transcription_metadata['language']
            
            cells = metadata_table.rows[3].cells
            cells[0].text = 'Durée'
            cells[1].text = self.format_duration(self.transcription_metadata['duration'])
            
            cells = metadata_table.rows[4].cells
            cells[0].text = 'Temps de traitement'
            cells[1].text = self.format_duration(self.transcription_metadata['processing_time'])
            
            doc.add_paragraph('')  # Ajouter un espace
            
            # Vérifier si nous avons des timestamps
            has_timestamps = self.transcription_metadata.get('has_timestamps', False)
            if has_timestamps and isinstance(self.transcription_result, dict) and "chunks" in self.transcription_result:
                doc.add_heading('Transcription avec horodatages', level=1)
                
                # Ajouter un tableau pour les horodatages
                timestamps_table = doc.add_table(rows=1, cols=2)
                timestamps_table.style = 'Table Grid'
                
                # En-têtes du tableau
                header_cells = timestamps_table.rows[0].cells
                header_cells[0].text = 'Temps'
                header_cells[1].text = 'Texte'
                
                # Ajouter les chunks avec horodatages
                for chunk in self.transcription_result["chunks"]:
                    start = chunk.get("timestamp", [0])[0]
                    # Formater le temps en minutes:secondes
                    minutes = int(start // 60)
                    seconds = int(start % 60)
                    time_str = f"{minutes:02d}:{seconds:02d}"
                    chunk_text = chunk.get("text", "")
                    
                    row_cells = timestamps_table.add_row().cells
                    row_cells[0].text = time_str
                    row_cells[1].text = chunk_text
            else:
                # Ajouter la transcription normale
                doc.add_heading('Transcription', level=1)
                
                # Ajouter le texte de la transcription
                if isinstance(self.transcription_result, dict) and "text" in self.transcription_result:
                    doc.add_paragraph(self.transcription_result["text"])
                else:
                    doc.add_paragraph(str(self.transcription_result))
            
            # Sauvegarder le document
            doc.save(save_path)
            
            self.status_var.set(f"Transcription sauvegardée: {save_path}")
            return True
        except Exception as e:
            self.status_var.set(f"Erreur lors de la sauvegarde: {str(e)}")
            messagebox.showerror("Erreur", f"Erreur lors de la sauvegarde: {str(e)}")
            return False
        
    def save_transcription(self):
        """Sauvegarder la transcription dans un fichier texte ou docx"""
        if not hasattr(self, '_transcription_result') or not self._transcription_result:
            messagebox.showerror("Erreur", "Aucune transcription à sauvegarder")
            return
            
        # Récupérer le format de sortie
        output_format = self.output_format.get()
        
        # Réglages de base pour le dialogue de sauvegarde
        save_options = {
            "initialdir": os.path.dirname(self.file_path.get()) if self.file_path.get() else os.path.expanduser("~"),
            "parent": self
        }
        
        # Adapter les options en fonction du format
        if output_format == "txt":
            save_options.update({
                "title": "Sauvegarder la transcription en TXT",
                "defaultextension": ".txt",
                "filetypes": [("Fichiers texte", "*.txt"), ("Tous les fichiers", "*.*")],
            })
        else:  # docx
            save_options.update({
                "title": "Sauvegarder la transcription en DOCX",
                "defaultextension": ".docx",
                "filetypes": [("Documents Word", "*.docx"), ("Tous les fichiers", "*.*")],
            })
        
        # Proposer un nom de fichier par défaut basé sur le fichier audio original
        if self.file_path.get():
            base_name = os.path.splitext(os.path.basename(self.file_path.get()))[0]
            save_options["initialfile"] = f"{base_name}_transcription"
        
        # Ouvrir le dialogue de sauvegarde
        save_path = filedialog.asksaveasfilename(**save_options)
        
        if not save_path:
            return  # L'utilisateur a annulé
            
        try:
            self.status_var.set("Sauvegarde en cours...")
            
            # Sauvegarder selon le format choisi
            if output_format == "txt":
                self.save_txt_file(save_path)
            else:  # docx
                self.save_docx_file(save_path)
                
            # Afficher un message de succès
            messagebox.showinfo("Succès", f"Transcription sauvegardée dans {save_path}")
            self.status_var.set(f"Sauvegardé dans {os.path.basename(save_path)}")
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la sauvegarde: {str(e)}")
            self.status_var.set("Erreur lors de la sauvegarde")

    def on_beam_change(self, value):
        """Callback lorsque l'option de qualité change - CORRIGÉ"""
        if value == "Rapide":
            self.beam_size_var.set(1)
        elif value == "Standard":
            self.beam_size_var.set(2)  # Réduit de 3 à 2
        else:  # Élevée
            self.beam_size_var.set(3)  # Réduit de 5 à 3
        self.save_preferences()

    def format_duration(self, seconds):
        """Formater une durée en secondes en format lisible"""
        if seconds < 60:
            return f"{seconds:.1f} secondes"
        else:
            minutes = int(seconds // 60)
            secs = int(seconds % 60)
            if minutes < 60:
                return f"{minutes} minute{'s' if minutes > 1 else ''} {secs} seconde{'s' if secs > 1 else ''}"
            else:
                hours = int(minutes // 60)
                mins = int(minutes % 60)
                return f"{hours} heure{'s' if hours > 1 else ''} {mins} minute{'s' if mins > 1 else ''} {secs} seconde{'s' if secs > 1 else ''}"


if __name__ == "__main__":
    try:
        # Check if accelerate is available and warn if not
        if not HAS_ACCELERATE:
            print("WARNING: The 'accelerate' package is not installed.")
            print("The application will still function, but for better performance")
            print("consider installing it with: pip install 'accelerate>=0.26.0'")
            
        app = ModernAudioTranscriptionApp()
        
        # Détecter le périphérique
        if torch.cuda.is_available():
            device_name = torch.cuda.get_device_name(0)
            app.device_info.set(f"GPU: {device_name}")
        else:
            app.device_info.set("Périphérique: CPU")
            
        app.mainloop()
    except Exception as e:
        print(f"Erreur lors du démarrage de l'application: {e}")
        import traceback
        traceback.print_exc()